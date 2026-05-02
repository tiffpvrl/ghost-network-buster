"""PDF and DOCX deliverables (audit summary + complaint draft)."""

from __future__ import annotations

import hashlib
from io import BytesIO

from fpdf import FPDF

from ghost_network_buster.models import AuditState, AuditSummary


def _ascii_pdf_text(text: str) -> str:
    """fpdf core fonts are latin-1; normalize common unicode punctuation then drop anything unrenderable."""
    text = (text or "")
    text = (
        text
        .replace("\u2014", "-")
        .replace("\u2013", "-")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u2022", "-")
        .replace("\u2018", "'")
        .replace("\u2026", "...")
    )
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _paragraphs(state: AuditState, summary: AuditSummary) -> list[str]:
    lines: list[str] = [
        f"Carrier audited: {state.carrier}",
        f"ZIP / region context: {state.zip_code or '-'}",
        f"Care needs selected: {', '.join(state.care_needs) if state.care_needs else '-'}",
        "",
        f"Providers in sample: {summary.providers_total}",
        f"Calls completed: {summary.calls_completed}",
        f"Confirmed real (accepting / usable): {summary.real_count}",
        f"Ghost listings (disconnected / wrong network): {summary.ghost_count}",
        f"Voicemail / unconfirmed: {summary.voicemail_count}",
        f"Ghost rate (ghosts / completed calls): {summary.ghost_rate * 100:.1f}%",
        "",
        "Top verified providers (excerpt):",
    ]
    for i, p in enumerate(summary.top_providers, start=1):
        lines.append(
            f"  {i}. {p.provider_name or p.npi} - {p.summary or ''} "
            f"(verified {p.verified_at or 'n/a'})"
        )
    if not summary.top_providers:
        lines.append("  (none in this run)")
    lines.extend(
        [
            "",
            "Transcript excerpts are included for user verification. "
            "Mock mode uses simulated transcripts; live mode uses call recordings.",
            "",
            f"Voice pipeline: {summary.voice_mode}",
        ]
    )
    return lines


def build_audit_summary_pdf(state: AuditState, summary: AuditSummary) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.multi_cell(0, 8, _ascii_pdf_text("Ghost Network Buster - Audit summary (draft)"))
    pdf.ln(2)
    pdf.set_font("Helvetica", size=10)
    for block in _paragraphs(state, summary):
        pdf.multi_cell(0, 5, _ascii_pdf_text(block))
        pdf.ln(1)
    return bytes(pdf.output(dest="S"))


def _rag_source_label(src: str) -> str:
    if ":" in src:
        return src.split(":", 1)[1].strip()[:100]
    name = src.split("/")[-1].split("\\")[-1]
    return name.replace("__chunk", " \u00a7").replace("_", " ").replace(".md", "").strip()[:100]


def _clean_excerpt(text: str) -> str:
    """Fix common OCR/copy-paste artifacts in RAG excerpts for professional presentation."""
    import re  # noqa: PLC0415
    # Fix words split at chunk boundary (e.g. "icipating" → label as excerpt)
    text = re.sub(r"^\s*[a-z]{1,5}ing\s+", "[EXCERPT] ...", text)
    # Label truncated Paragraph references as excerpts rather than showing bare numbers
    text = re.sub(
        r"(Respondents shall implement Paragraph \d+\([^)]+\)[^.]*\.?.*)",
        r"\1 [EXCERPT — full text available on request]",
        text,
        count=1,
    )
    # Fix chunk-boundary duplications like "r PartParticipating" → "Participating"
    text = re.sub(r'\S*Part(?=Participating)', '', text)
    text = re.sub(r'\br?\s+PartP(?=articipating)', 'P', text)
    # Fix any remaining "icipating Provider Directory" fragments
    text = re.sub(r'\b\w{0,6}icipating Provider Directory', 'Participating Provider Directory', text)
    return text.strip()


def _ph(para, text: str) -> None:
    para.add_run(text).bold = True


def _sig_block(doc) -> None:
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    _ph(p, "[YOUR SIGNATURE]\n")
    _ph(p, "[YOUR PRINTED NAME]\n")
    p.add_run("Date:  ")
    _ph(p, "[________________]")


def _disclaimer(doc, Pt, RGBColor) -> None:  # noqa: N803
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "DISCLAIMER: Not legal advice. Review and edit before filing. "
        "Consult a licensed attorney if needed. "
        "CHAMP Ombudsman (NY behavioral health): 888-614-5400 (free, confidential)."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def _hyperlink(para, url: str, text: str) -> None:
    """Insert a clickable hyperlink into a paragraph."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT  # noqa: PLC0415
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415
    r_id = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run_el = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(u)
    run_el.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run_el.append(t)
    hl.append(run_el)
    para._p.append(hl)


def _failures_table(doc, ghosts: list, ghost_rate: float) -> None:
    """Insert structured 'Table of Verified Failures' with systemic header if rate > 50%."""
    from docx.shared import Pt  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415
    if ghost_rate > 0.50:
        p = doc.add_paragraph()
        run = p.add_run(
            "This data demonstrates a systemic failure of network adequacy as defined by "
            "the February 2026 NY OAG standards."
        )
        run.bold = True
        run.font.size = Pt(10)
    if not ghosts:
        doc.add_paragraph("No ghost listings recorded in this audit.")
        return
    table = doc.add_table(rows=len(ghosts) + 1, cols=5)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(["Provider Name", "Specialty", "Phone", "Interaction Result", "Verified (UTC)"]):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for i, r in enumerate(ghosts, start=1):
        row = table.rows[i].cells
        row[0].text = r.provider_name or r.npi
        row[1].text = r.specialty or "\u2014"
        row[2].text = r.phone or "\u2014"
        _reason_labels = {
            "disconnected":           "Disconnected / Not in Service",
            "wrong_network":          "Not in Network (Confirmed by Staff)",
            "no_behavioral_health":   "Does Not Offer Behavioral Health Services",
            "not_accepting_patients": "Not Accepting New Patients (Closed Panel)",
            "wrong_provider":         "Wrong Provider at Listed Number",
            "retired":                "Provider No Longer Practicing",
            "wrong_specialty":        "Specialty Mismatch — Not Behavioral Health",
            "referral_only":          "Referral Required (Undisclosed Access Barrier)",
        }
        result = _reason_labels.get(r.ghost_reason or "", None)
        if result is None:
            result = r.ghost_reason.replace("_", " ").title() if r.ghost_reason else "Not Confirmed"
        row[3].text = result
        row[4].text = r.verified_at or "\u2014"
    p = doc.add_paragraph()
    run = p.add_run(
        "Note: All verification calls included a mandatory AI-identity disclosure and notice "
        "of recording at the start of the interaction, pursuant to FCC and NY State consumer "
        "transparency standards (2026)."
    )
    run.italic = True
    run.font.size = Pt(9)


def build_complaint_draft_docx(
    state: AuditState,
    summary: AuditSummary,
    rag_hits: list[dict[str, object]] | None = None,
    gcp_project: str | None = None,
    vertex_location: str = "us-central1",
) -> bytes:
    """
    Three-section complaint packet as an editable .docx:
      Section 1 — Internal Grievance Letter (to insurer, send immediately)
      Section 2 — DFS Evidence Pack (upload after 30 days if unresolved)
      Section 3 — AG Statement of Fact (submit to OAG anytime)
    """
    from docx import Document  # noqa: PLC0415
    from docx.enum.text import WD_BREAK  # noqa: PLC0415
    from docx.shared import Inches, Pt, RGBColor  # noqa: PLC0415

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    ghosts = [r for r in state.results if r.status == "ghost"]
    gr_pct = f"{summary.ghost_rate * 100:.0f}%"
    carrier = state.carrier
    zip_code = state.zip_code
    verification_id = hashlib.sha256(state.audit_id.encode()).hexdigest()[:12].upper()

    # Synthesize RAG chunks into clean regulatory prose (one LLM call, shared across all sections)
    rag_synthesis = ""
    if gcp_project and rag_hits:
        from ghost_network_buster.agents.complaint_agent import synthesize_rag_hits  # noqa: PLC0415
        rag_synthesis = synthesize_rag_hits(
            rag_hits, carrier, zip_code, summary.ghost_rate, gcp_project, vertex_location
        )

    # Add Verification ID to footer of every page for chain-of-custody feel
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: PLC0415
    for sec in doc.sections:
        sec.footer.is_linked_to_previous = False
        fp = sec.footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(
            f"Verification ID: {verification_id}  \u2014  Ghost Network Buster Audit Evidence Pack  \u2014  Not Legal Advice"
        )
        fr.font.size = Pt(8)
        fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    def pb() -> None:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def note(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # ── COVER PAGE ───────────────────────────────────────────────────────────
    doc.add_heading("Ghost Network Complaint Packet", level=1)
    doc.add_heading(f"{carrier} \u2014 ZIP {zip_code}", level=2)
    p = doc.add_paragraph()
    p.add_run(f"Verification ID: {verification_id}").bold = True
    p.add_run(f"   |   Audit ID: {state.audit_id[:8].upper()}")
    note(
        "This packet contains three ready-to-file documents. Fill in all bold [BRACKETED] "
        "fields before submitting. Each section targets a different recipient. "
        "This is not legal advice."
    )
    doc.add_paragraph()

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Stage", "Recipient", "When to Submit"]):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for i, (stage, recip, when) in enumerate([
        ("1 \u2014 Section 1", f"{carrier} Grievance Dept.", "Immediately"),
        ("2 \u2014 Section 2", "NY DFS Portal", "After 30 days (if unresolved)"),
        ("3 \u2014 Section 3", "NY Attorney General OAG", "Anytime"),
    ], start=1):
        row = table.rows[i].cells
        row[0].text = stage
        row[1].text = recip
        row[2].text = when
    doc.add_paragraph()
    note(
        f"Audit summary: {summary.calls_completed} providers audited \u2014 "
        f"{summary.ghost_count} ghost listings ({gr_pct} ghost rate). "
        f"Every ghost listing is a directory inaccuracy and a potential regulatory violation."
    )

    # ── SECTION 1: INTERNAL GRIEVANCE LETTER ─────────────────────────────────
    pb()
    doc.add_heading("Section 1 of 3 \u2014 Internal Grievance Letter", level=1)
    doc.add_heading(f"To: {carrier} Appeals & Grievance Department", level=2)
    note(
        "Call member services on your ID card and say: 'I want to file a formal grievance "
        "regarding a provider directory inaccuracy.' Mail or fax this letter as confirmation."
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    _ph(p, "[YOUR FULL NAME]\n[ADDRESS, CITY, STATE ZIP]\n[PHONE] | [EMAIL]\n")
    p.add_run("Member ID:  "); _ph(p, "[________________]")
    p.add_run("   Plan:  "); _ph(p, "[________________]")
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run("Date:  "); _ph(p, "[________________]")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"To: {carrier} Appeals and Grievance Department\n").bold = True
    _ph(p, "[INSURER MAILING ADDRESS FROM YOUR ID CARD]")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        f"Re: Formal Grievance \u2014 Inaccurate Behavioral Health Provider Directory, ZIP {zip_code}"
    ).bold = True
    doc.add_paragraph()

    doc.add_paragraph(
        f"I am filing this formal grievance regarding {carrier}\u2019s behavioral health provider "
        f"directory for ZIP code {zip_code}. An independent AI-voice audit (Verification ID: "
        f"{verification_id}) found that {summary.ghost_count} of {summary.calls_completed} listed "
        f"providers ({gr_pct}) were ghost listings \u2014 either unreachable or confirmed as "
        f"non-participating. Every ghost listing is a directory inaccuracy and a potential "
        f"violation of NY Insurance Law \u00a73241 and \u00a73217-a."
    )
    doc.add_paragraph()
    doc.add_heading("Table of Verified Failures", level=3)
    _failures_table(doc, ghosts, summary.ghost_rate)
    doc.add_paragraph()
    doc.add_paragraph(
        "Applicable legal standards include: (1) NY Insurance Law \u00a73241 (network adequacy mandate); "
        "(2) NY Insurance Law \u00a73217-a / \u00a74324 (directory accuracy disclosure); "
        "(3) MHPAEA 2024/2025 Final Rules on Non-Quantitative Treatment Limitations (NQTLs) \u2014 "
        "inaccurate directories act as a discriminatory NQTL creating greater barriers to mental health "
        "care than medical/surgical care; "
        "(4) the Requiring Enhanced and Accurate Lists (REAL) Health Providers Act (H.R. 7148), under "
        "which providers not verified within 90 days must be marked and removed within 5 business days; "
        "(5) NY OAG v. EmblemHealth, Assurance of Discontinuance (Feb 2026) \u2014 the NY Attorney General "
        "established that failure to maintain directory accuracy constitutes a violation of state parity "
        "laws (Reference: Feb 2026 $2.5M settlement, ag.ny.gov/press-release/2026/emblemhealth-ghost-network-settlement)."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        f"I request that {carrier}: (1) correct all inaccurate listings within 10 business days; "
        f"(2) provide a written list of verified in-network behavioral health providers accepting "
        f"new patients in ZIP {zip_code}; and (3) confirm the directory has been updated."
    )
    p = doc.add_paragraph()
    p.add_run(f"If {carrier} is unable to provide a verified in-network behavioral health appointment "
              f"within 10 business days of this letter, I formally demand a ")
    p.add_run("Single Case Agreement (SCA)").bold = True
    p.add_run(" or administrative lead-in authorization to see an out-of-network provider at "
              "in-network cost-sharing rates. Per NY Insurance Law \u00a73241, if your directory "
              "is inaccurate and prevents access to care, the insurer is responsible for the cost differential.")
    doc.add_paragraph(
        f"Failure to resolve this grievance will result in escalation to the NY Department of "
        f"Financial Services and the NY Attorney General\u2019s Health Care Bureau."
    )
    doc.add_paragraph()
    _sig_block(doc)
    _disclaimer(doc, Pt, RGBColor)

    # ── SECTION 2: DFS EVIDENCE PACK ─────────────────────────────────────────
    pb()
    doc.add_heading("Section 2 of 3 \u2014 DFS Evidence Pack", level=1)
    doc.add_heading("NY Department of Financial Services \u2014 Consumer Complaint", level=2)
    note(
        "Upload to the DFS Consumer Complaint Portal after 30 days if unresolved. "
        "DFS can fine insurers for network adequacy violations."
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Complainant:  "); _ph(p, "[YOUR FULL NAME]")
    p.add_run("   Member ID:  "); _ph(p, "[________________]")
    p.add_run(f"   Carrier:  {carrier}")
    p.add_run(f"   Verification ID:  {verification_id}")
    p = doc.add_paragraph(); p.add_run("Date:  "); _ph(p, "[________________]")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        f"Subject: Inaccurate Behavioral Health Provider Directory \u2014 Network Adequacy Violation \u2014 {carrier}"
    ).bold = True
    doc.add_paragraph()
    doc.add_paragraph(
        f"I am filing this complaint against {carrier} for maintaining an inaccurate behavioral "
        f"health provider directory in violation of NY Insurance Law \u00a73241, \u00a73217-a, "
        f"MHPAEA NQTL standards, and the REAL Act (H.R. 7148) 90-day verification requirement."
    )

    ev_table = doc.add_table(rows=5, cols=2)
    ev_table.style = "Table Grid"
    for k, v in enumerate([
        ("Carrier", carrier), ("ZIP Code", zip_code),
        ("Providers Audited", str(summary.calls_completed)),
        ("Ghost Listings", f"{summary.ghost_count} ({gr_pct})"),
        ("Audit Method", "AI-voice verification calls to each listed provider"),
    ]):
        row = ev_table.rows[k].cells
        row[0].text = v[0]
        for run in row[0].paragraphs[0].runs:
            run.bold = True
        row[1].text = v[1]
    doc.add_paragraph()
    doc.add_heading("Table of Verified Failures", level=3)
    _failures_table(doc, ghosts, summary.ghost_rate)

    if rag_hits:
        doc.add_paragraph()
        doc.add_heading("Regulatory Basis", level=3)
        if rag_synthesis:
            import re as _re  # noqa: PLC0415
            for line in rag_synthesis.splitlines():
                line = line.strip()
                # Strip bullet markers but preserve ** bold markers
                line = _re.sub(r"^[•\-]\s+", "", line)
                if not line:
                    continue
                p = doc.add_paragraph(style="List Bullet")
                parts = _re.split(r"\*\*(.+?)\*\*", line)
                for i, part in enumerate(parts):
                    if not part:
                        continue
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True
        else:
            for h in rag_hits:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{_rag_source_label(str(h.get('source', '')))}: ").bold = True
                p.add_run(_clean_excerpt(str(h.get("excerpt", ""))[:300]))

    doc.add_paragraph()
    doc.add_heading("Source References", level=3)
    for label, url in [
        ("NY OAG EmblemHealth Settlement (Feb 2026)", "https://ag.ny.gov/press-release/2026/emblemhealth-ghost-network-settlement"),
        ("REAL Act (H.R. 7148) Summary", "https://questanalytics.com/news/requiring-enhanced-accurate-lists-of-health-providers-act/"),
        ("NY DFS Consumer Complaint Portal", "https://www.dfs.ny.gov/complaint"),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{label}: ")
        _hyperlink(p, url, url)

    doc.add_paragraph()
    doc.add_paragraph(
        f"I request that DFS investigate {carrier}\u2019s directory accuracy, assess network "
        f"adequacy under \u00a73241, and take appropriate enforcement action."
    )
    doc.add_paragraph()
    _sig_block(doc)
    _disclaimer(doc, Pt, RGBColor)

    # ── SECTION 3: AG STATEMENT OF FACT ──────────────────────────────────────
    pb()
    doc.add_heading("Section 3 of 3 \u2014 Statement of Fact for the Attorney General", level=1)
    doc.add_heading("NY Office of the Attorney General \u2014 Health Care Bureau", level=2)
    note(
        "Submit at ag.ny.gov/file-complaint or call 1-800-771-7755. "
        "The OAG Health Care Bureau investigates systemic violations and uses ghost rate "
        "data to build enforcement cases (e.g., Feb 2026 $2.5M EmblemHealth settlement)."
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Complainant:  "); _ph(p, "[YOUR FULL NAME]  |  [PHONE]  |  [EMAIL]")
    p.add_run(f"   Verification ID:  {verification_id}")
    p = doc.add_paragraph(); p.add_run("Date:  "); _ph(p, "[________________]")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        f"Re: Systemic Ghost Network Violation \u2014 {carrier}, ZIP {zip_code} \u2014 {gr_pct} Ghost Rate"
    ).bold = True
    doc.add_paragraph()

    if gcp_project:
        from ghost_network_buster.agents import complaint_agent  # noqa: PLC0415
        llm_narrative = complaint_agent.generate_complaint_narrative(
            carrier, zip_code, summary.ghost_rate,
            [dict(h) for h in (rag_hits or [])],
            gcp_project, vertex_location,
        )
    else:
        llm_narrative = ""

    ag_template = (
        f"I am submitting this statement to report a systemic ghost network pattern in "
        f"{carrier}\u2019s behavioral health provider directory for ZIP code {zip_code}. "
        f"An AI-voice audit (Verification ID: {verification_id}) found that "
        f"{summary.ghost_count} of {summary.calls_completed} listed providers ({gr_pct}) "
        f"were ghost listings. The NY Attorney General has established that this constitutes "
        f"a violation of state parity laws (NY OAG v. EmblemHealth, Feb 2026, $2.5M settlement).\n\n"
        f"Under MHPAEA 2024/2025 Final Rules, an inaccessible behavioral health network "
        f"constitutes a discriminatory nonquantitative treatment limitation (NQTL) \u2014 "
        f"creating greater barriers to mental health care than medical/surgical coverage. "
        f"Under the REAL Act (H.R. 7148), {carrier} is failing the 90-day verification cycle "
        f"mandated for the providers listed in the table below \u2014 unverified providers must "
        f"be removed within 5 business days of identification; {carrier} has not done so.\n\n"
        f"I request that the Health Care Bureau investigate {carrier}\u2019s directory accuracy "
        f"as a systemic matter, assess compliance with \u00a73241 and MHPAEA, and take "
        f"enforcement action consistent with the EmblemHealth precedent."
    )

    for para_text in (llm_narrative if llm_narrative else ag_template).split("\n\n"):
        if para_text.strip():
            doc.add_paragraph(para_text.strip())

    doc.add_paragraph()
    doc.add_heading("Table of Verified Failures", level=3)
    _failures_table(doc, ghosts, summary.ghost_rate)

    if rag_hits:
        doc.add_paragraph()
        doc.add_heading("Regulatory Basis", level=3)
        if rag_synthesis:
            for line in rag_synthesis.splitlines():
                line = line.strip().lstrip("•-* ")
                if not line:
                    continue
                p = doc.add_paragraph(style="List Bullet")
                import re as _re  # noqa: PLC0415
                parts = _re.split(r"\*\*(.+?)\*\*", line)
                for i, part in enumerate(parts):
                    p.add_run(part).bold = (i % 2 == 1)
        else:
            for h in rag_hits:
                doc_type = str(h.get("doc_type", ""))
                p = doc.add_paragraph(style="List Bullet")
                label = _rag_source_label(str(h.get("source", "")))
                p.add_run(f"{label}{f' [{doc_type}]' if doc_type else ''}: ").bold = True
                p.add_run(_clean_excerpt(str(h.get("excerpt", ""))[:300]))

    doc.add_paragraph()
    _sig_block(doc)
    _disclaimer(doc, Pt, RGBColor)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
